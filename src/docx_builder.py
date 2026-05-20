from io import BytesIO
from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
import re

TEMPLATE_PATH = "templates/Vergo_Report_Template.docx"

PLACEHOLDER_KEYS = {
    "TASK_NAME",
    "CLIENT_NAME",
    "SITE_NAME",
    "ASSESSMENT_DATE",
    "ASSESSMENT_METHOD",
    "VIDEO_DURATION",
    "ASSESSOR",
    "ASSESSMENT_OVERVIEW",
    "SUMMARY_OF_RESULTS",
    "RISK_EXPOSURE_ANALYSIS",
    "OVERALL_OBSERVATIONS",
    "RECOMMENDATIONS",
    "TRAINING_VIDEOS",
    "DISCLAIMER",
}

SECTION_PLACEHOLDERS = {
    "{{ASSESSMENT_OVERVIEW}}",
    "{{SUMMARY_OF_RESULTS}}",
    "{{RISK_EXPOSURE_ANALYSIS}}",
    "{{OVERALL_OBSERVATIONS}}",
    "{{RECOMMENDATIONS}}",
    "{{TRAINING_VIDEOS}}",
    "{{DISCLAIMER}}",
}

SECTION_HEADING_PREFIXES = [
    "Section 1 – Assessment Overview",
    "Section 1 - Assessment Overview",
    "Assessment Overview",
    "Section 2 – Summary of Assessment Results",
    "Section 2 - Summary of Assessment Results",
    "Summary of Assessment Results",
    "Section 3 – Task-Based Risk Exposure Analysis",
    "Section 3 - Task-Based Risk Exposure Analysis",
    "Task-Based Risk Exposure Analysis",
    "Section 4 – Overall Observations",
    "Section 4 - Overall Observations",
    "Overall Observations",
    "Section 5 – Overall Recommendations",
    "Section 5 - Overall Recommendations",
    "Overall Recommendations",
    "Section 6 – Targeted Vergo Training Videos",
    "Section 6 - Targeted Vergo Training Videos",
    "Targeted Vergo Training Videos",
    "Section 7 – Disclaimer",
    "Section 7 - Disclaimer",
    "Disclaimer",
]


def _normalize_text(value):
    if value is None:
        return ""
    if isinstance(value, str):
        return value.strip()
    if isinstance(value, (int, float)):
        return str(value)
    if isinstance(value, list):
        parts = []
        for item in value:
            text = _normalize_text(item)
            if text:
                parts.append(text)
        return "\n\n".join(parts)
    if isinstance(value, dict):
        parts = []
        for key, val in value.items():
            text = _normalize_text(val)
            if text:
                parts.append(f"{key}: {text}")
        return "\n".join(parts)
    return str(value).strip()


def _strip_section_heading_prefix(text):
    if not isinstance(text, str):
        return text
    normalized = text.strip()
    for prefix in SECTION_HEADING_PREFIXES:
        if normalized.startswith(prefix):
            remainder = normalized[len(prefix):].lstrip(':–—- ').strip()
            return remainder
    return normalized


def _normalize_section_value(value):
    if value is None:
        return []
    if isinstance(value, list):
        return value
    return [value]


def _build_assessment_overview(value):
    items = []
    paragraphs = []
    if isinstance(value, dict):
        paragraphs = value.get("paragraphs", []) or []
    elif isinstance(value, list):
        paragraphs = value
    elif isinstance(value, str):
        paragraphs = [value]
    else:
        paragraphs = [str(value)]

    for paragraph in paragraphs:
        text = _normalize_text(paragraph)
        if text:
            items.append({"text": _strip_section_heading_prefix(text), "bold": False, "bullet": False})
    return items


def _build_summary_of_results(value):
    items = []
    paragraphs = []
    key_exposures = []

    if isinstance(value, dict):
        paragraphs = value.get("paragraphs", []) or []
        if isinstance(paragraphs, str):
            paragraphs = [paragraphs]

        key_exposures = value.get("key_exposures", []) or []
        if isinstance(key_exposures, str):
            key_exposures = [key_exposures]

        if not paragraphs:
            lines = []
            overall_score = value.get("overall_score")
            risk_level = value.get("risk_level")
            num_frames = value.get("num_frames_analyzed") or value.get("frames_analyzed")
            main_exposure = value.get("main_exposure_pattern")
            interpretation = value.get("interpretation")

            if overall_score not in (None, ""):
                lines.append(f"Overall score: {overall_score}.")
            if risk_level not in (None, ""):
                lines.append(f"Risk level: {risk_level}.")
            if num_frames not in (None, ""):
                lines.append(f"Frames analyzed: {num_frames}.")
            if main_exposure not in (None, ""):
                lines.append(f"Main exposure pattern: {main_exposure}.")
            if interpretation not in (None, ""):
                lines.append(interpretation)

            if lines:
                paragraphs = [" ".join(lines)]

        if not key_exposures and value.get("distribution_summary"):
            distribution = value.get("distribution_summary") or []
            for distribution_item in distribution:
                if isinstance(distribution_item, dict):
                    label = distribution_item.get("band_label") or distribution_item.get("score_range")
                    percent = distribution_item.get("frames_percent")
                    if label and percent is not None:
                        key_exposures.append(f"{label}: {percent}%")
                else:
                    item_text = _normalize_text(distribution_item)
                    if item_text:
                        key_exposures.append(item_text)

    elif isinstance(value, list):
        for item in value:
            if isinstance(item, str):
                paragraphs.append(item)
            elif isinstance(item, dict):
                heading = item.get("heading") or item.get("title")
                if heading:
                    paragraphs.append(str(heading))
                body = item.get("body") or item.get("details") or item.get("text")
                if body:
                    paragraphs.append(_normalize_text(body))
            else:
                paragraphs.append(str(item))
    elif isinstance(value, str):
        paragraphs = [value]
    else:
        paragraphs = [str(value)]

    for paragraph in paragraphs:
        text = _normalize_text(paragraph)
        if text:
            items.append({"text": _strip_section_heading_prefix(text), "bold": False, "bullet": False})

    for exposure in key_exposures:
        text = _normalize_text(exposure)
        if text:
            items.append({"text": _strip_section_heading_prefix(text), "bold": False, "bullet": True})

    return items


def _build_risk_exposure_analysis(value):
    items = []
    normalized = _normalize_section_value(value)
    for item in normalized:
        if isinstance(item, str):
            text = _strip_section_heading_prefix(item)
            if text:
                items.append({"text": text, "bold": False, "bullet": False})
        elif isinstance(item, dict):
            heading = item.get("heading") or item.get("title") or item.get("section")
            body = item.get("body") or item.get("paragraphs") or item.get("details") or item.get("text")
            if heading:
                heading_text = _strip_section_heading_prefix(_normalize_text(heading))
                if heading_text:
                    items.append({"text": heading_text, "bold": True, "bullet": False})
            if isinstance(body, list):
                body_text = " ".join([_normalize_text(p) for p in body if _normalize_text(p)])
            else:
                body_text = _normalize_text(body)
            if body_text:
                items.append({"text": _strip_section_heading_prefix(body_text), "bold": False, "bullet": False})
        else:
            text = _strip_section_heading_prefix(_normalize_text(item))
            if text:
                items.append({"text": text, "bold": False, "bullet": False})
    return items


def _build_overall_observations(value):
    items = []
    paragraphs = []
    bullets = []

    if isinstance(value, dict):
        paragraphs = value.get("paragraphs", []) or []
        bullets = value.get("bullets", []) or []
        if isinstance(paragraphs, str):
            paragraphs = [paragraphs]
        if isinstance(bullets, str):
            bullets = [bullets]
    elif isinstance(value, list):
        for item in value:
            if isinstance(item, str):
                paragraphs.append(item)
            elif isinstance(item, dict):
                if item.get("heading") and item.get("body"):
                    paragraphs.append(f"{item.get('heading')}: {item.get('body')}")
                else:
                    bullets.append(item.get("text") or item.get("bullet") or str(item))
            else:
                paragraphs.append(str(item))
    elif isinstance(value, str):
        paragraphs = [value]
    else:
        paragraphs = [str(value)]

    for paragraph in paragraphs:
        text = _normalize_text(paragraph)
        if text:
            items.append({"text": _strip_section_heading_prefix(text), "bold": False, "bullet": False})

    for bullet in bullets:
        text = _normalize_text(bullet)
        if text:
            items.append({"text": _strip_section_heading_prefix(text), "bold": False, "bullet": True})

    return items


def _build_recommendations(value):
    items = []
    normalized = _normalize_section_value(value)

    for item in normalized:
        if isinstance(item, str):
            text = _strip_section_heading_prefix(item)
            if text:
                items.append({"text": text, "bold": False, "bullet": False})
        elif isinstance(item, dict):
            heading = item.get("heading") or item.get("title") or item.get("recommendation")
            body = item.get("body") or item.get("details") or item.get("text")
            if heading:
                heading_text = _strip_section_heading_prefix(_normalize_text(heading))
                if heading_text:
                    items.append({"text": heading_text, "bold": True, "bullet": False})
            if isinstance(body, list):
                body_text = " ".join([_normalize_text(p) for p in body if _normalize_text(p)])
            else:
                body_text = _normalize_text(body)
            if body_text:
                items.append({"text": _strip_section_heading_prefix(body_text), "bold": False, "bullet": False})
        else:
            text = _strip_section_heading_prefix(_normalize_text(item))
            if text:
                items.append({"text": text, "bold": False, "bullet": False})

    return items


def _build_training_videos(value):
    items = []
    normalized = _normalize_section_value(value)

    for item in normalized:
        if isinstance(item, str):
            text = _normalize_text(item)
            if text:
                items.append({"text": _strip_section_heading_prefix(text), "bold": True, "bullet": False})
        elif isinstance(item, dict):
            module = item.get("module") or item.get("name") or item.get("video")
            reason = item.get("reason") or item.get("rationale") or item.get("why") or item.get("details")
            module_text = _normalize_text(module)
            reason_text = _normalize_text(reason)
            if module_text:
                items.append({"text": _strip_section_heading_prefix(module_text), "bold": True, "bullet": False})
            if reason_text:
                items.append({"text": _strip_section_heading_prefix(reason_text), "bold": False, "bullet": False})
        else:
            text = _normalize_text(item)
            if text:
                items.append({"text": _strip_section_heading_prefix(text), "bold": True, "bullet": False})

    return items


def _build_disclaimer(value):
    text = _normalize_text(value)
    text = _strip_section_heading_prefix(text)
    return [{"text": text, "bold": False, "bullet": False}] if text else []


def _render_placeholder_value(placeholder, value):
    if placeholder == "{{ASSESSMENT_OVERVIEW}}":
        return _build_assessment_overview(value)
    if placeholder == "{{SUMMARY_OF_RESULTS}}":
        return _build_summary_of_results(value)
    if placeholder == "{{RISK_EXPOSURE_ANALYSIS}}":
        return _build_risk_exposure_analysis(value)
    if placeholder == "{{OVERALL_OBSERVATIONS}}":
        return _build_overall_observations(value)
    if placeholder == "{{RECOMMENDATIONS}}":
        return _build_recommendations(value)
    if placeholder == "{{TRAINING_VIDEOS}}":
        return _build_training_videos(value)
    if placeholder == "{{DISCLAIMER}}":
        return _build_disclaimer(value)
    return [{"text": _normalize_text(value), "bold": False, "bullet": False}]


def _clear_paragraph(paragraph):
    for run in paragraph.runs[:]:
        paragraph._element.remove(run._element)


def _replace_paragraph_text(paragraph, text):
    style = paragraph.style
    _clear_paragraph(paragraph)
    if text:
        run = paragraph.add_run(text)
        run.bold = False
    paragraph.style = style
    return paragraph


def _insert_paragraph_after(paragraph, text=None, style=None):
    new_p = OxmlElement('w:p')
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style is not None:
        new_para.style = style
    if text:
        run = new_para.add_run(text)
        run.bold = False
    return new_para


def _replace_paragraph_with_items(paragraph, placeholder, items):
    if not items:
        paragraph._element.getparent().remove(paragraph._element)
        return True

    full_text = paragraph.text
    style = paragraph.style

    if full_text.strip() == placeholder:
        first = items[0]
        _replace_paragraph_text(paragraph, ("• " if first["bullet"] else "") + first["text"])
        if first["bold"] and paragraph.runs:
            paragraph.runs[0].bold = True

        last_paragraph = paragraph
        for item in items[1:]:
            text_value = ("• " if item["bullet"] else "") + item["text"]
            new_para = _insert_paragraph_after(last_paragraph, text_value, style=style)
            if item["bold"] and new_para.runs:
                new_para.runs[0].bold = True
            last_paragraph = new_para
        return True

    replacement_text = "\n\n".join([("• " if item["bullet"] else "") + item["text"] for item in items])
    new_text = full_text.replace(placeholder, replacement_text)
    _replace_paragraph_text(paragraph, new_text)
    return True


def _replace_text_in_paragraph(paragraph, placeholder, replacement):
    if placeholder not in paragraph.text:
        return False

    if placeholder in SECTION_PLACEHOLDERS:
        items = _render_placeholder_value(placeholder, replacement)
        return _replace_paragraph_with_items(paragraph, placeholder, items)

    replacement_text = _normalize_text(replacement)
    new_text = paragraph.text.replace(placeholder, replacement_text)
    _replace_paragraph_text(paragraph, new_text)
    return True


def _iter_table_paragraphs(table):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                yield paragraph
            for nested_table in cell.tables:
                yield from _iter_table_paragraphs(nested_table)


def _iter_all_paragraphs(document):
    for paragraph in document.paragraphs:
        yield paragraph
    for table in document.tables:
        yield from _iter_table_paragraphs(table)
    for section in document.sections:
        for paragraph in section.header.paragraphs:
            yield paragraph
        for table in section.header.tables:
            yield from _iter_table_paragraphs(table)
        for paragraph in section.footer.paragraphs:
            yield paragraph
        for table in section.footer.tables:
            yield from _iter_table_paragraphs(table)


def _find_placeholders_in_text(text):
    if not text:
        return []
    return re.findall(r"\{\{[^}]+\}\}", text)


def _find_placeholders_in_document(document):
    placeholders = set()
    for paragraph in _iter_all_paragraphs(document):
        placeholders.update(_find_placeholders_in_text(paragraph.text))
    return placeholders


def _replace_placeholders(document, placeholder_values):
    expected = set(placeholder_values.keys())
    found = _find_placeholders_in_document(document)

    missing_from_template = expected - found
    extra_in_template = found - expected

    print(f"DEBUG: placeholders found in template before replacement: {sorted(found)}")
    print(f"DEBUG: placeholders expected: {sorted(expected)}")
    if missing_from_template:
        print(f"DEBUG: placeholders expected but missing from template: {sorted(missing_from_template)}")
    if extra_in_template:
        print(f"DEBUG: extra placeholders found in template not provided by code: {sorted(extra_in_template)}")

    replaced_placeholders = set()

    for placeholder, replacement in placeholder_values.items():
        placeholder_found = False
        for paragraph in list(_iter_all_paragraphs(document)):
            if _replace_text_in_paragraph(paragraph, placeholder, replacement):
                placeholder_found = True
                replaced_placeholders.add(placeholder)
        if not placeholder_found:
            print(f"WARNING: placeholder {placeholder} was not found in the template.")

    print(f"DEBUG: placeholders replaced: {sorted(replaced_placeholders)}")

    remaining = _find_placeholders_in_document(document)
    if remaining:
        print(f"ERROR: placeholders remaining after replacement: {sorted(remaining)}")
        raise ValueError(f"Unreplaced placeholders remain in document: {sorted(remaining)}")


def _get_placeholder_values(report):
    cover = report.get("cover_details", {}) or {}
    return {
        "{{TASK_NAME}}": _normalize_text(cover.get("task")),
        "{{CLIENT_NAME}}": _normalize_text(cover.get("company")),
        "{{SITE_NAME}}": _normalize_text(cover.get("site_location") or cover.get("site_name")),
        "{{ASSESSMENT_DATE}}": _normalize_text(cover.get("assessment_date")),
        "{{ASSESSMENT_METHOD}}": _normalize_text(cover.get("assessment_method")),
        "{{VIDEO_DURATION}}": _normalize_text(cover.get("video_duration")),
        "{{ASSESSOR}}": _normalize_text(cover.get("assessor")),
        "{{ASSESSMENT_OVERVIEW}}": report.get("assessment_overview", []),
        "{{SUMMARY_OF_RESULTS}}": report.get("score_summary", {}),
        "{{RISK_EXPOSURE_ANALYSIS}}": report.get("risk_exposure_analysis", []),
        "{{OVERALL_OBSERVATIONS}}": report.get("overall_observations", {}),
        "{{RECOMMENDATIONS}}": report.get("recommendations", []),
        "{{TRAINING_VIDEOS}}": report.get("training_videos", []),
        "{{DISCLAIMER}}": report.get("disclaimer", ""),
    }


def build_docx(report: dict) -> bytes:
    document = Document(TEMPLATE_PATH)
    placeholder_values = _get_placeholder_values(report)
    _replace_placeholders(document, placeholder_values)

    output = BytesIO()
    document.save(output)
    output.seek(0)
    return output.read()
