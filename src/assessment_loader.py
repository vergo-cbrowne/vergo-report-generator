from typing import Any, Dict, List
from io import BytesIO

from googleapiclient.http import MediaIoBaseDownload
from docx import Document

import google_drive


TEXT_MIME_TYPES = {
    "text/plain",
    "text/markdown",
    "application/octet-stream",
}

DOCX_MIME_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
GOOGLE_DOC_MIME_TYPE = "application/vnd.google-apps.document"

TASK_NOTE_KEYWORDS = ("task", "note", "notes", "context")


def _find_report_json(files: List[Dict[str, Any]]) -> Dict[str, Any] | None:
    normalized_target = "report.json"
    exact_match = next((item for item in files if item["name"] == normalized_target), None)
    if exact_match:
        return exact_match

    case_insensitive_matches = [
        item for item in files if item["name"].lower() == normalized_target.lower()
    ]
    return case_insensitive_matches[0] if case_insensitive_matches else None


def _print_files(title: str, files: List[Dict[str, Any]]) -> None:
    print(title)
    for item in files:
        print(
            f"- name={item.get('name')} mimeType={item.get('mimeType')} id={item.get('id')} "
            f"parents={item.get('parents')} driveId={item.get('driveId')}"
        )


def load_assessment_folder(service, folder_id: str) -> tuple[Dict[str, Any], List[Dict[str, Any]]]:
    files = google_drive.list_folder_files(service, folder_id)

    _print_files("Assessment folder contents:", files)
    report_file = _find_report_json(files)

    if report_file is None:
        folder_metadata = google_drive.get_file_metadata(
            service,
            folder_id,
            fields="id,name,mimeType,driveId,parents",
        )
        print(
            f"Selected folder metadata: id={folder_metadata.get('id')} name={folder_metadata.get('name')} "
            f"mimeType={folder_metadata.get('mimeType')} driveId={folder_metadata.get('driveId')}"
        )

        print("Fallback: searching globally for report.json files")
        global_report_files = google_drive.search_files_by_name_global(service, "report.json")

        if not global_report_files:
            raise FileNotFoundError(
                "report.json not found in assessment folder and no global report.json files were discovered"
            )

        print("Global report.json candidates:")
        for item in global_report_files:
            parent_ids = item.get("parents") or []
            parent_id = parent_ids[0] if parent_ids else "<none>"
            parent_name = (
                google_drive.get_folder_name(service, parent_id) if parent_ids else "<unknown>"
            )
            print(
                f"- fileId={item.get('id')} parentId={parent_id} parentName={parent_name} "
                f"driveId={item.get('driveId')}"
            )

        direct_matches = [
            item for item in global_report_files if folder_id in (item.get("parents") or [])
        ]
        if direct_matches:
            report_file = direct_matches[0]
            print("Found report.json with selected folder as parent.")
        else:
            fallback_candidate = global_report_files[0]
            fallback_parents = fallback_candidate.get("parents") or []
            if not fallback_parents:
                raise FileNotFoundError(
                    "report.json found globally, but no parent folder IDs were available for fallback"
                )

            fallback_parent_id = fallback_parents[0]
            fallback_parent_name = google_drive.get_folder_name(service, fallback_parent_id)
            print(
                f"No direct report.json parent match. Using fallback parent folder: "
                f"id={fallback_parent_id} name={fallback_parent_name}"
            )
            files = google_drive.list_folder_files(service, fallback_parent_id)
            _print_files("Fallback folder contents:", files)
            report_file = fallback_candidate
            folder_id = fallback_parent_id

    report_json = google_drive.download_json_file(service, report_file["id"])
    snapshot_files = [item for item in files if "snapshot" in item["name"].lower()]

    return report_json, snapshot_files


def _download_drive_file_bytes(service, file_id: str) -> bytes:
    request = service.files().get_media(
        fileId=file_id,
        supportsAllDrives=True,
    )

    file_buffer = BytesIO()
    downloader = MediaIoBaseDownload(file_buffer, request)

    done = False
    while not done:
        _, done = downloader.next_chunk()

    return file_buffer.getvalue()


def _export_google_doc_text(service, file_id: str) -> str:
    request = service.files().export_media(
        fileId=file_id,
        mimeType="text/plain",
    )

    file_buffer = BytesIO()
    downloader = MediaIoBaseDownload(file_buffer, request)

    done = False
    while not done:
        _, done = downloader.next_chunk()

    return file_buffer.getvalue().decode("utf-8", errors="replace").strip()


def _decode_text_bytes(raw_bytes: bytes) -> str:
    for encoding in ("utf-8", "utf-8-sig", "latin-1"):
        try:
            return raw_bytes.decode(encoding, errors="replace").strip()
        except Exception:
            continue

    return raw_bytes.decode("utf-8", errors="replace").strip()


def _extract_docx_text(raw_bytes: bytes) -> str:
    document = Document(BytesIO(raw_bytes))

    parts: list[str] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)

    for table in document.tables:
        for row in table.rows:
            row_values = []
            for cell in row.cells:
                cell_text = " ".join(
                    paragraph.text.strip()
                    for paragraph in cell.paragraphs
                    if paragraph.text.strip()
                )
                if cell_text:
                    row_values.append(cell_text)

            if row_values:
                parts.append(" | ".join(row_values))

    return "\n".join(parts).strip()


def _is_task_note_candidate(file_item: Dict[str, Any]) -> bool:
    name = (file_item.get("name") or "").strip()
    lower_name = name.lower()
    mime_type = file_item.get("mimeType") or ""

    if not name:
        return False

    if lower_name == "task_context.txt":
        return True

    if lower_name.endswith(".txt"):
        return True

    if lower_name.endswith(".md"):
        return True

    if lower_name.endswith(".docx"):
        return True

    if mime_type in TEXT_MIME_TYPES and lower_name != "report.json":
        return True

    if mime_type == DOCX_MIME_TYPE:
        return True

    if mime_type == GOOGLE_DOC_MIME_TYPE:
        return any(keyword in lower_name for keyword in TASK_NOTE_KEYWORDS)

    return False


def _truncate_text(text: str, max_chars: int = 6000) -> str:
    text = (text or "").strip()

    if len(text) <= max_chars:
        return text

    return text[:max_chars].rstrip() + "\n\n[Note truncated for prompt length.]"


def load_task_notes(service, folder_id: str) -> Dict[str, Any]:
    """
    Load lightweight task notes from the assessment folder.

    Supported:
    - task_context.txt
    - any .txt / .md file
    - any .docx file
    - Google Docs files where the name includes task, notes, or context

    These notes are passed to the AI as context. They are not rendered directly
    into the report.
    """
    try:
        files = google_drive.list_folder_files(service, folder_id)
    except Exception as exc:
        print(f"Task notes: Could not list folder files: {exc}")
        return {"combined_text": "", "files": []}

    candidates = [item for item in files if _is_task_note_candidate(item)]

    # Avoid loading report/debug/status files as notes.
    candidates = [
        item for item in candidates
        if (item.get("name") or "").strip().lower()
        not in {"report.json", "status.json", "baseline_summary.csv"}
    ]

    if not candidates:
        print("Task notes: No task note files found.")
        return {"combined_text": "", "files": []}

    loaded_files: list[dict[str, str]] = []

    for item in candidates:
        file_id = item.get("id")
        name = item.get("name") or "Unnamed note file"
        mime_type = item.get("mimeType") or ""

        if not file_id:
            continue

        try:
            lower_name = name.lower()

            if mime_type == GOOGLE_DOC_MIME_TYPE:
                text = _export_google_doc_text(service, file_id)
            elif mime_type == DOCX_MIME_TYPE or lower_name.endswith(".docx"):
                raw_bytes = _download_drive_file_bytes(service, file_id)
                text = _extract_docx_text(raw_bytes)
            else:
                raw_bytes = _download_drive_file_bytes(service, file_id)
                text = _decode_text_bytes(raw_bytes)

            text = _truncate_text(text)

            if not text:
                print(f"Task notes: Skipped empty note file: {name}")
                continue

            loaded_files.append(
                {
                    "name": name,
                    "mimeType": mime_type,
                    "text": text,
                }
            )

            print(f"Task notes: Loaded note file: {name}")

        except Exception as exc:
            print(f"Task notes: Could not load note file {name}: {exc}")

    if not loaded_files:
        print("Task notes: Candidate note files were found, but none could be loaded.")
        return {"combined_text": "", "files": []}

    combined_parts = []
    for file_item in loaded_files:
        combined_parts.append(
            f"NOTE FILE: {file_item['name']}\n{file_item['text']}"
        )

    combined_text = "\n\n---\n\n".join(combined_parts)

    print(f"Task notes: Loaded {len(loaded_files)} note file(s).")
    print("Task notes: Notes will be passed to OpenAI as contextual evidence.")

    return {
        "combined_text": _truncate_text(combined_text, max_chars=12000),
        "files": loaded_files,
    }

def load_local_assessment_folder(folder_path):
    """Load an extracted local assessment folder containing report.json and snapshots."""
    from pathlib import Path
    import json

    folder = Path(folder_path)

    report_json = folder / "report.json"
    if not report_json.exists():
        matches = list(folder.rglob("report.json"))
        if not matches:
            raise FileNotFoundError(f"No report.json found in {folder}")
        report_json = matches[0]
        folder = report_json.parent

    report_data = json.loads(report_json.read_text())

    snapshot_files = []
    for pattern in ["*.png", "*.jpg", "*.jpeg"]:
        snapshot_files.extend(folder.rglob(pattern))

    snapshot_files = [
        str(p) for p in snapshot_files
        if "__MACOSX" not in str(p)
    ]

    return report_data, snapshot_files
